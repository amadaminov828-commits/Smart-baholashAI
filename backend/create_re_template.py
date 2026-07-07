"""
Ko'chmas mulk baholash hisoboti uchun professional DOCX shablon yaratish.
Ishga tushirish: .\venv\Scripts\python.exe create_re_template.py
"""

import os
import django

os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'valuation_platform.settings')
django.setup()

from docx import Document
from docx.shared import Pt, Cm, Mm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MEDIA_ROOT = 'c:/Users/Asus/Desktop/antigravity/backend/media'
OUTPUT_PATH = os.path.join(MEDIA_ROOT, 'templates', 'KochmasMulk_shablon.docx')

def set_font(run, name='Times New Roman', size=12, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def para(doc, text='', align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=False, 
         color=None, space_before=0, space_after=6, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = RGBColor(*color)
    return p

def divider(doc, color=(0, 70, 140)):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '%02X%02X%02X' % color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def two_col_table(doc, rows_data):
    """Label | Value juftligi uchun jadval"""
    table = doc.add_table(rows=len(rows_data), cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(rows_data):
        row = table.rows[i]
        lc = row.cells[0]
        vc = row.cells[1]
        lc.width = Cm(7)
        vc.width = Cm(11)
        lp = lc.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        lr = lp.add_run(label)
        set_font(lr, bold=True, size=11)
        vp = vc.paragraphs[0]
        vr = vp.add_run(str(value))
        set_font(vr, size=11)
    doc.add_paragraph()

def create_template():
    doc = Document()

    # Sahifa sozlamalari (A4)
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    # ─── HEADER (Sarlavha) ───────────────────────────────────────────────────
    header = section.header
    htable = header.add_table(1, 3, width=Cm(17))
    htable.style = 'Table Grid'

    # Logo (chap)
    lc = htable.rows[0].cells[0]
    lc.width = Cm(3)
    logo_p = lc.paragraphs[0]
    logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_run = logo_p.add_run()
    logo_run.add_text('{{ logo }}')
    set_font(logo_run, size=9, italic=True, color=(100,100,100))

    # Kompaniya nomi (o'rta)
    mc = htable.rows[0].cells[1]
    mc.width = Cm(10)
    mp = mc.paragraphs[0]
    mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = mp.add_run('MULK BAHOSI EKSPERT MARKAZI')
    set_font(mr, bold=True, size=11, color=(0, 51, 102))
    mp2 = mc.add_paragraph('Baholash va Konsalting Xizmatlari')
    mp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(mp2.runs[0] if mp2.runs else mp2.add_run(''), size=9, italic=True, color=(80,80,80))

    # Raqam (o'ng)
    rc = htable.rows[0].cells[2]
    rc.width = Cm(4)
    rp = rc.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rr = rp.add_run('Hisobot: №{{ ID }}')
    set_font(rr, size=9, color=(100,100,100))
    rp2 = rc.add_paragraph('{{ valuation_date }}')
    rp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(rp2.runs[0] if rp2.runs else rp2.add_run(''), size=9)

    # ─── BODY ────────────────────────────────────────────────────────────────

    # Sarlavha
    p_title = para(doc, "KO'CHMAS MULK OBYEKTINI BAHOLASH HISOBOTI",
                   align=WD_ALIGN_PARAGRAPH.CENTER, size=16, bold=True,
                   color=(0, 51, 102), space_before=6, space_after=4)

    p_sub = para(doc, "Baholash maqsadi: {{ PURPOSE }}",
                 align=WD_ALIGN_PARAGRAPH.CENTER, size=12, bold=False,
                 color=(80, 80, 80), space_after=8)

    divider(doc, color=(0, 70, 140))

    # ─── I. Umumiy ma'lumotlar ───────────────────────────────────────────────
    para(doc, "I. UMUMIY MA'LUMOTLAR", size=13, bold=True, color=(0, 51, 102), space_before=8)
    divider(doc)

    two_col_table(doc, [
        ("Hisobot raqami",       "{{ ID }}"),
        ("Baholash sanasi",      "{{ valuation_date }}"),
        ("Shartnoma raqami",     "{{ SHARTNOMA }}"),
        ("Shartnoma sanasi",     "{{ SANA }}"),
        ("Baholash maqsadi",     "{{ PURPOSE }}"),
    ])

    # ─── II. Mulkdor ma'lumotlari ────────────────────────────────────────────
    para(doc, "II. MULKDOR MA'LUMOTLARI", size=13, bold=True, color=(0, 51, 102), space_before=8)
    divider(doc)

    two_col_table(doc, [
        ("F.I.Sh.",               "{{ FISH }}"),
        ("Pasport seriyasi",      "{{ PASPORT }}"),
        ("JSHSHIR",               "{{ JSHSHIR }}"),
        ("Pasport berilgan joy",  "{{ BERILDI }}"),
        ("Viloyat",               "{{ VILOYAT }}"),
        ("Tuman",                 "{{ TUMAN }}"),
    ])

    # ─── III. Mulk tavsifi ──────────────────────────────────────────────────
    para(doc, "III. MULK TAVSIFI", size=13, bold=True, color=(0, 51, 102), space_before=8)
    divider(doc)

    two_col_table(doc, [
        ("Kadastr raqami",   "{{ KADASTR }}"),
        ("Manzil",           "{{ MANZIL }}"),
        ("Qurilish yili",    "{{ YEAR }}"),
        ("Umumiy maydon",    "{{ AREA }} m²"),
        ("Mulk turi",        "Turar-joy (yakka tartibdagi)"),
    ])

    # ─── IV. Baholash natijalari ─────────────────────────────────────────────
    para(doc, "IV. BAHOLASH NATIJALARI", size=13, bold=True, color=(0, 51, 102), space_before=10)
    divider(doc)

    para(doc, "Quyidagi uch yondashuv asosida bozor qiymati aniqlandi:", size=11, italic=True, space_after=4)

    # Yondashuvlar jadvali
    approach_table = doc.add_table(rows=5, cols=3)
    approach_table.style = 'Table Grid'
    approach_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['Yondashuv', 'Hisoblangan qiymat (so\'m)', 'Izoh']
    for j, h in enumerate(headers):
        cell = approach_table.rows[0].cells[j]
        run = cell.paragraphs[0].add_run(h)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(run, bold=True, size=11, color=(255,255,255))
        # Ko'k fon
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '00468C')
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)

    rows_data = [
        ('Xarajat yondashuvi',    '{{ cost_final_val }}',  'Tiklanish qiymati asosida'),
        ('Daromad yondashuvi',    '{{ inc_final_val }}',   'Kapitallashtirish asosida'),
        ('Qiyosiy yondashuv',     '{{ comp_final_val }}',  'Analog narxlar asosida'),
        ('YAKUNIY BOZOR QIYMATI','{{ final_market_value }}','Tanlangan yondashuv'),
    ]
    for i, (y, v, iz) in enumerate(rows_data, 1):
        r = approach_table.rows[i]
        r.cells[0].paragraphs[0].add_run(y)
        r.cells[1].paragraphs[0].add_run(v)
        r.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r.cells[2].paragraphs[0].add_run(iz)
        set_font(r.cells[0].paragraphs[0].runs[0], size=11, bold=(i==4))
        set_font(r.cells[1].paragraphs[0].runs[0], size=11, bold=(i==4), color=(0,100,0) if i==4 else None)
        set_font(r.cells[2].paragraphs[0].runs[0], size=10, italic=True)
        if i == 4:
            for c in r.cells:
                tcPr = c._tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), 'E8F5E9')
                shd.set(qn('w:val'), 'clear')
                tcPr.append(shd)

    doc.add_paragraph()

    # Yakuniy qiymat so'z bilan
    p_final = para(doc, "Yakuniy bozor qiymati so'z bilan:", size=12, bold=True, space_before=6, space_after=2)
    p_words = para(doc, "{{ final_value_words }}", size=12, bold=False, italic=True, color=(0,100,0), space_after=10)

    divider(doc, color=(0, 100, 0))

    # ─── V. Garov uchun maxsus bo'lim ───────────────────────────────────────
    para(doc, "V. GAROV QIYMATI TO'G'RISIDA XULOSA", size=13, bold=True, color=(0, 51, 102), space_before=10)
    divider(doc)

    para(doc,
         "Ushbu hisobot O'zbekiston Respublikasining ko'chmas mulkni baholash bo'yicha "
         "amaldagi qonunchiligiga muvofiq tuzilgan. Mulkning bozor qiymati yuqorida "
         "ko'rsatilgan miqdorda aniqlanib, bank garoviga qo'yish uchun asos bo'lib xizmat qiladi.",
         size=11, space_after=4)

    para(doc, "Baholash ushbu ko'rsatkichlarga asoslanadi:", size=11, bold=True, space_after=2)
    for item in [
        "✓  Xarajat yondashuvi: mulkni qayta tiklash xarajatlari asosida",
        "✓  Daromad yondashuvi: mulkdan olinadigan ijara daromadlari asosida",
        "✓  Qiyosiy yondashuv: o'xshash mulklar bozor narxlari asosida",
    ]:
        p_item = doc.add_paragraph(style='List Bullet')
        p_item.paragraph_format.space_after = Pt(2)
        r_item = p_item.runs[0] if p_item.runs else p_item.add_run(item)
        r_item.text = item
        set_font(r_item, size=11)

    doc.add_paragraph()

    # ─── VI. Imzolar ────────────────────────────────────────────────────────
    para(doc, "VI. TASDIQLASH VA IMZOLAR", size=13, bold=True, color=(0, 51, 102), space_before=10)
    divider(doc)

    sign_table = doc.add_table(rows=4, cols=2)
    sign_table.style = 'Table Grid'

    sign_data = [
        ("Baholovchi:", "{{ approver_name }}"),
        ("Imzo:",       "{% if show_seal %}{{ signature }}{% endif %}"),
        ("Mudur muhri:","{% if show_seal %}{{ seal }}{% endif %}"),
        ("Sana:",       "{{ valuation_date }}"),
    ]
    for i, (lbl, val) in enumerate(sign_data):
        sign_table.rows[i].cells[0].paragraphs[0].add_run(lbl)
        set_font(sign_table.rows[i].cells[0].paragraphs[0].runs[-1], bold=True, size=11)
        sign_table.rows[i].cells[1].paragraphs[0].add_run(val)
        set_font(sign_table.rows[i].cells[1].paragraphs[0].runs[-1], size=11)

    doc.add_paragraph()

    # QR kod
    para(doc, "Hisobot haqiqiyligini tekshirish:", size=10, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=2)
    p_qr = para(doc, '', align=WD_ALIGN_PARAGRAPH.CENTER)
    p_qr.add_run('{{ qr_code }}')

    para(doc, "Ushbu QR kodni skanerlang va hisobot haqiqiyligini tekshiring",
         size=9, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=(100,100,100))

    # ─── FOOTER ─────────────────────────────────────────────────────────────
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("Hisobot №{{ ID }} | {{ valuation_date }} | Ushbu hujjat elektron tarzda tasdiqlangan")
    set_font(fr, size=8, italic=True, color=(120,120,120))

    # Saqlash
    doc.save(OUTPUT_PATH)
    print(f"✅ Shablon yaratildi: {OUTPUT_PATH}")

    # DB ga ro'yxatdan o'tkazish
    from reports.models import ReportTemplate
    t, created = ReportTemplate.objects.get_or_create(
        name="Ko'chmas mulk - Professional shablon",
        object_type='real_estate'
    )
    t.file = 'templates/KochmasMulk_shablon.docx'
    t.is_default = True
    t.save()

    # Eski defaultlarni o'chirish
    ReportTemplate.objects.filter(
        object_type='real_estate', is_default=True
    ).exclude(id=t.id).update(is_default=False)

    print(f"✅ DB ga qo'shildi: ID {t.id} | {'Yangi' if created else 'Yangilandi'}")
    print(f"\nBarcha real_estate shablonlari:")
    for tmpl in ReportTemplate.objects.filter(object_type='real_estate'):
        print(f"  - ID: {tmpl.id} | {tmpl.name} | default: {tmpl.is_default}")

if __name__ == '__main__':
    create_template()
