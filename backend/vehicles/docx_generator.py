import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Color blue
    c = docx.oxml.shared.OxmlElement('w:color')
    c.set(docx.oxml.shared.qn('w:val'), "0000EE")
    rPr.append(c)

    # Underline
    u = docx.oxml.shared.OxmlElement('w:u')
    u.set(docx.oxml.shared.qn('w:val'), 'single')
    rPr.append(u)
    
    # Force Arial 12pt directly on the hyperlink run
    rFonts = docx.oxml.shared.OxmlElement('w:rFonts')
    rFonts.set(docx.oxml.shared.qn('w:ascii'), 'Arial')
    rFonts.set(docx.oxml.shared.qn('w:hAnsi'), 'Arial')
    rFonts.set(docx.oxml.shared.qn('w:cs'), 'Arial')
    rPr.append(rFonts)
    
    sz = docx.oxml.shared.OxmlElement('w:sz')
    sz.set(docx.oxml.shared.qn('w:val'), '24') # 12pt = 24 half-points
    rPr.append(sz)
    
    new_run.append(rPr)

    # Add text node
    t = docx.oxml.shared.OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)
    return hyperlink

def set_cell_arial_12(cell, bold=False):
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if not paragraph.runs:
            run = paragraph.add_run()
        for run in paragraph.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            run.bold = bold

def set_paragraph_arial_12(paragraph):
    if not paragraph.runs:
        return
    for run in paragraph.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(12)

from docx.text.paragraph import Paragraph

def iter_paragraphs(doc):
    # Main document body (includes tables, text boxes, shapes)
    for p in doc._element.xpath('.//w:p'):
        yield Paragraph(p, doc)
        
    # Headers and Footers
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header:
                for p in header._element.xpath('.//w:p'):
                    yield Paragraph(p, header)
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer:
                for p in footer._element.xpath('.//w:p'):
                    yield Paragraph(p, footer)

from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm

def fill_docx_template(template_path, output_path, context, qr_code_path=None):
    """
    Upgraded version using docxtpl for Jinja2 support ({% if %}, etc.)
    """
    temp_dir = os.path.dirname(output_path)
    report_status = context.get('report_status', 'draft')
    is_approved = (report_status == 'approved')
    stamp_date = context.get('{hisobot_sanasi}', datetime.now().strftime('%d.%m.%Y'))
    
    # Ensure verification QR code is generated and resolved
    verify_qr_file = qr_code_path
    if not (verify_qr_file and os.path.exists(verify_qr_file)):
        import qrcode
        verify_url = context.get('verify_absolute_url', 'http://smartbaholash.uz')
        verify_qr_file = os.path.join(temp_dir, "verify_qr.png")
        qr = qrcode.QRCode(version=1, box_size=3, border=1)
        qr.add_data(verify_url)
        qr.make(fit=True)
        img = qr.make_image(fill_color="black", back_color="white")
        img.save(verify_qr_file)
        
    # 1. Generate/resolve appraiser signature path
    sig_path = context.get('signature_path')
    if is_approved and not (sig_path and os.path.exists(sig_path)):
        from PIL import Image, ImageDraw, ImageFont
        img_sig = Image.new('RGBA', (500, 160), (255, 255, 255, 0))
        draw_sig = ImageDraw.Draw(img_sig)
        blue_sig = (12, 43, 128, 255) # Rich dark blue pen ink
        
        font_paths = ['C:\\Windows\\Fonts\\segoesc.ttf', 'C:\\Windows\\Fonts\\lhandw.ttf', 'C:\\Windows\\Fonts\\ariali.ttf']
        font_sig = None
        for fp in font_paths:
            try:
                font_sig = ImageFont.truetype(fp, 44)
                break
            except IOError:
                continue
        if font_sig is None:
            font_sig = ImageFont.load_default()
            
        appraiser_name = context.get('appraiser_name', 'B.M.Mirzabekov')
        
        # Draw text with slight rotation for realistic handwriting look
        text_layer = Image.new('RGBA', (500, 160), (255, 255, 255, 0))
        draw_text = ImageDraw.Draw(text_layer)
        draw_text.text((250, 65), appraiser_name, fill=blue_sig, font=font_sig, anchor='mm')
        
        rotated_text = text_layer.rotate(-2, resample=Image.BICUBIC, expand=False)
        img_sig.alpha_composite(rotated_text)
        
        # Draw realistic signature pen flourish
        flourish_pts = [
            (60, 110), (120, 114), (180, 116), (240, 115), 
            (300, 112), (360, 107), (410, 97), (430, 88), 
            (435, 83), (420, 102), (380, 122), (330, 131), 
            (290, 128), (280, 120), (320, 108), (380, 95), (450, 90)
        ]
        draw_sig.line(flourish_pts, fill=blue_sig, width=3)
        
        img_sig = img_sig.rotate(3, resample=Image.BICUBIC, expand=False)
        sig_path = os.path.join(temp_dir, 'digital_sig.png')
        img_sig.save(sig_path)
        context['signature_path'] = sig_path
        
    # 2. Generate/resolve company seal path
    seal_path = context.get('seal_path')
    if is_approved and not (seal_path and os.path.exists(seal_path)):
        from PIL import Image, ImageDraw, ImageFont
        import math
        
        def draw_arc_text(image, text, center, radius, start_angle, font, fill, direction=1):
            draw = ImageDraw.Draw(image)
            cx, cy = center
            chars = list(text)
            if not chars:
                return
            widths = []
            for c in chars:
                try:
                    w = draw.textlength(c, font=font)
                except AttributeError:
                    w = font.getsize(c)[0] if hasattr(font, 'getsize') else 16
                widths.append(w)
            total_w = sum(widths)
            span_angle = total_w / radius
            start_rad = math.radians(start_angle) - (span_angle / 2) * direction
            current_angle = start_rad
            for idx, c in enumerate(chars):
                w = widths[idx]
                char_angle = current_angle + (w / 2 / radius) * direction
                x = cx + radius * math.cos(char_angle)
                y = cy + radius * math.sin(char_angle)
                rot_deg = math.degrees(char_angle) + 90
                if direction == -1:
                    rot_deg += 180
                char_w = int(w * 2) or 24
                char_h = int(font.size * 2) or 24
                char_img = Image.new('RGBA', (char_w, char_h), (255, 255, 255, 0))
                char_draw = ImageDraw.Draw(char_img)
                char_draw.text((char_w // 2, char_h // 2), c, fill=fill, font=font, anchor='mm')
                rotated_char = char_img.rotate(-rot_deg, resample=Image.BICUBIC, expand=True)
                rw, rh = rotated_char.size
                px = int(x - rw / 2)
                py = int(y - rh / 2)
                image.alpha_composite(rotated_char, (px, py))
                current_angle += (w / radius) * direction
                
        img_seal = Image.new('RGBA', (600, 600), (255, 255, 255, 0))
        draw_seal = ImageDraw.Draw(img_seal)
        blue_seal = (15, 60, 150, 255) # Premium official blue
        
        # Concentric outer double circles
        draw_seal.ellipse([20, 20, 580, 580], outline=blue_seal, width=8)
        draw_seal.ellipse([45, 45, 555, 555], outline=blue_seal, width=3)
        
        # Center horizontal dividers
        draw_seal.line([120, 190, 480, 190], fill=blue_seal, width=2)
        draw_seal.line([120, 410, 480, 410], fill=blue_seal, width=2)
        
        try:
            font_main = ImageFont.truetype('C:\\Windows\\Fonts\\arial.ttf', 30)
            font_bold = ImageFont.truetype('C:\\Windows\\Fonts\\arialbd.ttf', 38)
            font_small = ImageFont.truetype('C:\\Windows\\Fonts\\arial.ttf', 24)
            font_arc = ImageFont.truetype('C:\\Windows\\Fonts\\arialbd.ttf', 26)
        except IOError:
            font_main = ImageFont.load_default()
            font_bold = ImageFont.load_default()
            font_small = ImageFont.load_default()
            font_arc = ImageFont.load_default()
            
        # Draw curved text along circular arcs
        draw_arc_text(img_seal, "KONSALTING XIZMATI", (300, 300), 215, -90, font_arc, blue_seal, 1)
        draw_arc_text(img_seal, "QO'QON SHAHRI", (300, 300), 215, 90, font_arc, blue_seal, -1)
        
        # Draw center content
        company_name = context.get('company_name', 'PERCEPTION VALUE')
        words = company_name.split()
        if len(words) >= 2:
            w1 = " ".join(words[:-1])
            w2 = words[-1]
            draw_seal.text((300, 230), w1, fill=blue_seal, font=font_bold, anchor='mm')
            draw_seal.text((300, 275), w2, fill=blue_seal, font=font_bold, anchor='mm')
        else:
            draw_seal.text((300, 250), company_name, fill=blue_seal, font=font_bold, anchor='mm')
            
        company_stir = context.get('company_stir', '301234567')
        draw_seal.text((300, 320), f"STIR: {company_stir}", fill=blue_seal, font=font_main, anchor='mm')
        draw_seal.text((300, 358), "TASDIQLANDI", fill=blue_seal, font=font_bold, anchor='mm')
        draw_seal.text((300, 390), stamp_date, fill=blue_seal, font=font_main, anchor='mm')
        
        # Slight rotation for realistic stamp placement
        img_seal = img_seal.rotate(-4, resample=Image.BICUBIC, expand=False)
        seal_path = os.path.join(temp_dir, 'digital_seal.png')
        img_seal.save(seal_path)
        context['seal_path'] = seal_path

    doc = DocxTemplate(template_path)
    
    # Process special image fields if they exist in context
    image_fields = ['logo_path', 'signature_path', 'seal_path']
    for field in image_fields:
        if field in context and context[field] and os.path.exists(context[field]):
            # Create InlineImage object and replace the path with the object
            tag_name = field.replace('_path', '')
            width = 30 if 'logo' in tag_name else 40
            context[tag_name] = InlineImage(doc, context[field], width=Mm(width))

    # Process QR code if provided
    if qr_code_path and os.path.exists(qr_code_path):
        # We can add it to context so user can use {{ qr_code }}, {{ verify_qr }} or other variants
        qr_image = InlineImage(doc, qr_code_path, width=Mm(25))
        context['qr_code'] = qr_image
        context['qr'] = qr_image
        context['verify_qr'] = qr_image
        context['verification_qr'] = qr_image
        
    # Set signature/seal aliases in context
    if 'signature' in context:
        context['appraiser_signature'] = context['signature']
        context['sig'] = context['signature']
        context['imzo'] = context['signature']
    if 'seal' in context:
        context['company_seal'] = context['seal']
        context['stamp'] = context['seal']
        context['pechat'] = context['seal']
    
    # Create a Jinja-friendly context where keys like '{tag}' become 'tag'
    jinja_context = {}
    for k, v in context.items():
        if isinstance(k, str) and k.startswith('{') and k.endswith('}'):
            jinja_context[k[1:-1]] = v
        jinja_context[k] = v
        
    # Render the template with jinja2 context
    doc.render(jinja_context)
    doc.save(output_path)
    
    # Reload using python-docx to avoid docxtpl clearing our programmatic changes during save
    doc_docx = docx.Document(output_path)
    
    # Ensure Table of Contents (MUNDARIJA) starts on its own page
    for p_idx, p in enumerate(doc_docx.paragraphs):
        p_text_clean = p.text.replace(" ", "").lower()
        if p_text_clean == "mundarija" or p_text_clean == "mundarija.":
            # Check if there is already a page break in the preceding paragraphs
            # to avoid duplicate page breaks creating a blank page
            has_break = False
            try:
                p_list = list(doc_docx.paragraphs)
                # Scan backwards from the paragraph before MUNDARIJA to find any page break
                for idx in range(p_idx - 1, -1, -1):
                    prev_p = p_list[idx]
                    xml_val = prev_p._p.xml
                    if 'w:br' in xml_val and 'page' in xml_val:
                        has_break = True
                        break
                    # If we find actual non-empty body text, we stop searching backwards
                    # because page breaks after content are what we care about.
                    if prev_p.text.strip():
                        break
            except Exception as break_err:
                print(f"Error checking break: {break_err}")
                
            # Commented out to prevent creating a blank Page 2. Templates naturally flow
            # Mundarija to page 2 since cover page contents fill page 1.
            # if not has_break:
            #     p.insert_paragraph_before().add_run().add_break(docx.enum.text.WD_BREAK.PAGE)
            break
    
    # EXACT string replacement for legacy single-brace tags like {hisobot_raqami}
    # Filter only string/number context values to replace in raw text
    placeholders_to_unbold = {
        '{makroiqtisodiy_jadval}', '{aholi_daromadlari_tahlili}', '{bozor_tendensiyalari_tahlili}',
        '{model_xususiyatlari_tahlili}', '{texnik_holat_tahlili}', '{bozor_vaziyati_matni}',
        '{iqtisodiy_vaziyat_matni}', '{muvofiqlashtirish_matni}', '{qiymat_turi_tahlili}',
        '{eskirish_matni}', '{daromad_yondashuvi_tavsifi}', '{qiyosiy_yondashuv_tavsifi}',
        '{xarajat_yondashuvi_tavsifi}'
    }

    # EXACT string replacement for legacy single-brace tags like {hisobot_raqami}
    # Filter only string/number context values to replace in raw text
    text_replacements = {}
    for k, v in context.items():
        if isinstance(k, str) and k.startswith('{') and k.endswith('}') and not isinstance(v, InlineImage):
            val_str = str(v)
            if k in placeholders_to_unbold or len(val_str) > 150:
                paragraphs = val_str.split('\n')
                formatted_paragraphs = []
                for p in paragraphs:
                    p = p.strip()
                    if p:
                        if not p.startswith('•') and not p.startswith('\t'):
                            p = '\t' + p
                        formatted_paragraphs.append(p)
                val_str = '\n'.join(formatted_paragraphs)
            text_replacements[k] = val_str

    # Sort replacements by key length descending to prevent partial match bugs
    sorted_reps = sorted(text_replacements.items(), key=lambda x: len(x[0]), reverse=True)

    from docx.oxml.shared import OxmlElement
    from docx.oxml.ns import qn

    def enforce_arial_normal(r_elem):
        rPrs = r_elem.xpath('./w:rPr')
        if not rPrs:
            rPr = OxmlElement('w:rPr')
            r_elem.insert(0, rPr)
        else:
            rPr = rPrs[0]
            
        for b in rPr.xpath('./w:b | ./w:bCs'):
            rPr.remove(b)
            
        # Explicitly set bold elements to false to prevent inheritance from bold paragraph styles
        b = OxmlElement('w:b')
        b.set(qn('w:val'), 'false')
        rPr.append(b)
        
        bCs = OxmlElement('w:bCs')
        bCs.set(qn('w:val'), 'false')
        rPr.append(bCs)
            
        for f in rPr.xpath('./w:rFonts'):
            rPr.remove(f)
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Arial')
        rFonts.set(qn('w:hAnsi'), 'Arial')
        rFonts.set(qn('w:cs'), 'Arial')
        rPr.append(rFonts)
        
        for sz in rPr.xpath('./w:sz | ./w:szCs'):
            rPr.remove(sz)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '24')
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), '24')
        rPr.append(szCs)

    def enforce_paragraph_unbold(p_elem):
        pPrs = p_elem.xpath('./w:pPr')
        if not pPrs:
            pPr = OxmlElement('w:pPr')
            p_elem.insert(0, pPr)
        else:
            pPr = pPrs[0]
            
        # Reset paragraph style to Normal to remove collapse triangle and reset heading spacing
        for pStyle in pPr.xpath('./w:pStyle'):
            pPr.remove(pStyle)
        pStyle = OxmlElement('w:pStyle')
        pStyle.set(qn('w:val'), 'Normal')
        pPr.append(pStyle)
        
        # Force Justified alignment (jc w:val="both")
        for jc in pPr.xpath('./w:jc'):
            pPr.remove(jc)
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'both')
        pPr.append(jc)
        
        # Force First-line indent (ind w:firstLine="709") - 709 dxa = 1.25 cm
        for ind in pPr.xpath('./w:ind'):
            pPr.remove(ind)
        ind = OxmlElement('w:ind')
        ind.set(qn('w:firstLine'), '709')
        pPr.append(ind)
        
        # Force Spacing: before=0, after=120 (6pt), line=276 (1.15 line spacing)
        for spacing in pPr.xpath('./w:spacing'):
            pPr.remove(spacing)
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:before'), '0')
        spacing.set(qn('w:after'), '120')
        spacing.set(qn('w:line'), '276')
        spacing.set(qn('w:lineRule'), 'auto')
        pPr.append(spacing)
            
        p_rPrs = pPr.xpath('./w:rPr')
        if not p_rPrs:
            p_rPr = OxmlElement('w:rPr')
            pPr.append(p_rPr)
        else:
            p_rPr = p_rPrs[0]
            
        for b in p_rPr.xpath('./w:b | ./w:bCs'):
            p_rPr.remove(b)
            
        b = OxmlElement('w:b')
        b.set(qn('w:val'), 'false')
        p_rPr.append(b)
        
        bCs = OxmlElement('w:bCs')
        bCs.set(qn('w:val'), 'false')
        p_rPr.append(bCs)

    def populate_formatted_text(p_elem, t_node, line_text, ns_prefix):
        # Reset text node
        t_node.text = ""
        r_elem = t_node.getparent()
        p_elem = r_elem.getparent()
        
        # We will insert new run elements right after r_elem, keeping the order!
        r_idx = p_elem.index(r_elem)
        insert_offset = 1
        
        parts = line_text.split('\t')
        for part_idx, part in enumerate(parts):
            if part_idx > 0:
                tab_run = OxmlElement(f'{ns_prefix}:r')
                tab = OxmlElement(f'{ns_prefix}:tab')
                tab_run.append(tab)
                p_elem.insert(r_idx + insert_offset, tab_run)
                insert_offset += 1
            
            if part:
                import re
                sub_parts = re.split(r'\*\*(.*?)\*\*', part)
                for sub_idx, sub_part in enumerate(sub_parts):
                    if not sub_part:
                        continue
                    
                    t_el = OxmlElement(f'{ns_prefix}:t')
                    t_el.set(qn('xml:space'), 'preserve')
                    t_el.text = sub_part
                    
                    if sub_idx % 2 == 1:
                        # Bold run
                        sub_r = OxmlElement(f'{ns_prefix}:r')
                        sub_rPr = OxmlElement(f'{ns_prefix}:rPr')
                        b = OxmlElement(f'{ns_prefix}:b')
                        sub_rPr.append(b)
                        bCs = OxmlElement(f'{ns_prefix}:bCs')
                        sub_rPr.append(bCs)
                        
                        rPr_exist = r_elem.xpath(f'./{ns_prefix}:rPr')
                        if rPr_exist:
                            for font in rPr_exist[0].xpath(f'./{ns_prefix}:rFonts'):
                                from copy import deepcopy
                                sub_rPr.append(deepcopy(font))
                            for sz in rPr_exist[0].xpath(f'./{ns_prefix}:sz'):
                                from copy import deepcopy
                                sub_rPr.append(deepcopy(sz))
                        sub_r.append(sub_rPr)
                        sub_r.append(t_el)
                        p_elem.insert(r_idx + insert_offset, sub_r)
                        insert_offset += 1
                    else:
                        if insert_offset > 1:
                            sub_r = OxmlElement(f'{ns_prefix}:r')
                            rPr_exist = r_elem.xpath(f'./{ns_prefix}:rPr')
                            if rPr_exist:
                                from copy import deepcopy
                                sub_rPr = deepcopy(rPr_exist[0])
                                for b in sub_rPr.xpath(f'./{ns_prefix}:b | ./{ns_prefix}:bCs'):
                                    sub_rPr.remove(b)
                                sub_r.append(sub_rPr)
                            sub_r.append(t_el)
                            p_elem.insert(r_idx + insert_offset, sub_r)
                            insert_offset += 1
                        else:
                            r_elem.append(t_el)

    def process_xml_paragraph(p_elem, ns_prefix):
        t_tab_xpath = f'.//{ns_prefix}:t | .//{ns_prefix}:tab'
        nodes = p_elem.xpath(t_tab_xpath)
        if not nodes: return
        
        parts = []
        for node in nodes:
            tag_local = node.tag.split('}')[-1]
            if tag_local == 't':
                if node.text:
                    parts.append(node.text)
            elif tag_local == 'tab':
                parts.append('\t')
                
        text = "".join(parts)
        if not text.strip(): return
        
        t_nodes = [n for n in nodes if n.tag.split('}')[-1] == 't']
        if not t_nodes: return
        
        is_header_paragraph = '{hisobot_sanasi}' in text and '{buyurtmachi}' in text and '\t' in text
        
        changed = False
        ai_generated = False
        force_bold = False
        
        placeholders_to_unbold = {
            '{makroiqtisodiy_jadval}', '{aholi_daromadlari_tahlili}', '{bozor_tendensiyalari_tahlili}',
            '{model_xususiyatlari_tahlili}', '{texnik_holat_tahlili}', '{bozor_vaziyati_matni}',
            '{iqtisodiy_vaziyat_matni}', '{muvofiqlashtirish_matni}', '{qiymat_turi_tahlili}',
            '{eskirish_matni}', '{daromad_yondashuvi_tavsifi}', '{qiyosiy_yondashuv_tavsifi}',
            '{xarajat_yondashuvi_tavsifi}'
        }
        
        for old_k, str_v in sorted_reps:
            if old_k in text:
                if old_k == '{hisobot_raqami}':
                    text = text.replace(old_k, f"**{str_v}**")
                else:
                    text = text.replace(old_k, str_v)
                changed = True
                if len(str_v) > 150 or old_k in placeholders_to_unbold:
                    ai_generated = True
                
        if changed:
            if is_header_paragraph:
                text = text.replace('..', '.')
                import re
                text = re.sub(r'\t+', '\t', text)
                text = re.sub(r'[ ]+', ' ', text)
                text = text.strip()
                if '\t' in text:
                    sub_parts = text.split('\t')
                    if len(sub_parts) >= 2 and sub_parts[1].strip() and not sub_parts[1].startswith('**'):
                        text = f"{sub_parts[0]}\t**{sub_parts[1].strip()}**"
                
            # Clear text in all text nodes
            for t in t_nodes: 
                t.text = ""
            
            # Remove all tab nodes in the paragraph so we can reconstruct them
            for node in nodes:
                tag_local = node.tag.split('}')[-1]
                if tag_local == 'tab':
                    parent = node.getparent()
                    if parent is not None:
                        parent.remove(node)
            
            if '\n' in text:
                # We have multiple paragraphs!
                # Split text by newline
                lines = [line.strip() for line in text.split('\n') if line.strip()]
                
                # Make deepcopies of the original unmodified p_elem for subsequent paragraphs
                copied_paragraphs = []
                for line in lines[1:]:
                    from copy import deepcopy
                    copied_paragraphs.append((deepcopy(p_elem), line))
                
                # First paragraph text is set on the current p_elem
                r_elem = t_nodes[0].getparent()
                for child in list(r_elem):
                    tag_local = child.tag.split('}')[-1]
                    if tag_local in ('t', 'tab', 'br') and child != t_nodes[0]:
                        r_elem.remove(child)
                
                populate_formatted_text(p_elem, t_nodes[0], lines[0], ns_prefix)
                if ai_generated and ns_prefix == 'w':
                    enforce_paragraph_unbold(p_elem)
                    for r in p_elem.xpath('.//w:r'):
                        enforce_arial_normal(r)
                
                # For the remaining paragraphs, insert and format them sequentially
                current_p = p_elem
                for new_p, line in copied_paragraphs:
                    new_t_nodes = new_p.xpath(f'.//{ns_prefix}:t')
                    if new_t_nodes:
                        new_t = new_t_nodes[0]
                        # Clear other runs/text in new_p
                        new_r = new_t.getparent()
                        for child in list(new_r):
                            tag_local = child.tag.split('}')[-1]
                            if tag_local in ('t', 'tab', 'br') and child != new_t:
                                new_r.remove(child)
                        
                        populate_formatted_text(new_p, new_t, line, ns_prefix)
                        if ai_generated and ns_prefix == 'w':
                            enforce_paragraph_unbold(new_p)
                            for r in new_p.xpath('.//w:r'):
                                enforce_arial_normal(r)
                                
                    current_p.addnext(new_p)
                    current_p = new_p
            else:
                if '\t' in text or '**' in text:
                    r_elem = t_nodes[0].getparent()
                    for child in list(r_elem):
                        tag_local = child.tag.split('}')[-1]
                        if tag_local in ('t', 'tab', 'br') and child != t_nodes[0]:
                            r_elem.remove(child)
                    populate_formatted_text(p_elem, t_nodes[0], text, ns_prefix)
                else:
                    t_nodes[0].text = text
            
            if force_bold and ns_prefix == 'w':
                r_elem = t_nodes[0].getparent()
                rPrs = r_elem.xpath('./w:rPr')
                if not rPrs:
                    rPr = OxmlElement('w:rPr')
                    r_elem.insert(0, rPr)
                else:
                    rPr = rPrs[0]
                if not rPr.xpath('./w:b'):
                    b = OxmlElement('w:b')
                    rPr.append(b)
                    
            if ai_generated and ns_prefix == 'w':
                enforce_paragraph_unbold(p_elem)
                for r in p_elem.xpath('.//w:r'):
                    enforce_arial_normal(r)

    # Apply brute force replace to all w:p and a:p in main document, headers, and footers
    parts = [doc_docx.part]
    
    # Prevent tables from autofitting and expanding beyond margins
    for table in doc_docx.tables:
        table.autofit = False
        table.allow_autofit = False
    for rel in doc_docx.part.rels.values():
        if "header" in rel.reltype or "footer" in rel.reltype:
            parts.append(rel.target_part)
            
    for part in parts:
        # Table row removal for requested tags
        for tr in part.element.xpath('.//w:tr'):
            tr_text = "".join(t.text for t in tr.xpath('.//w:t') if t.text)
            if "Dastlabki balans qiymati" in tr_text or "Balans qoldiq qiymati" in tr_text:
                parent = tr.getparent()
                if parent is not None:
                    parent.remove(tr)
                    
        for p in part.element.xpath('.//w:p'):
            p_text = "".join(t.text for t in p.xpath('.//w:t') if t.text)
            p_text_lower = p_text.lower()
            
            # Detect and clear the watermark/footer "Xisobotning X - varog'i"
            if "xisobotning" in p_text_lower and "varog" in p_text_lower:
                for t in p.xpath('.//w:t'):
                    t.text = ""
                continue
                
            process_xml_paragraph(p, 'w')
            
            # Post-process specific paragraphs that user wants unbolded but are static in template
            should_unbold = False
            
            if "ikki tomonlama tuzilgan shartnoma" in p_text_lower:
                should_unbold = True
            elif "baholashning aniq maqsadi" in p_text_lower:
                should_unbold = True
            elif "baholash ob'yektining nomi:" in p_text_lower or "baholash ob’yektining nomi:" in p_text_lower:
                should_unbold = True
            elif "huquq egasi (mulkdor):" in p_text_lower:
                should_unbold = True
            elif "ob'yektning tarkibi:" in p_text_lower or "ob’yektning tarkibi:" in p_text_lower:
                should_unbold = True
            elif "balansida asosiy vosita sifatida" in p_text_lower and ("ro‘yxatdan" in p_text_lower or "ro'yxatdan" in p_text_lower):
                should_unbold = True
            elif "garovdan yoki xatlovdan xoli" in p_text_lower:
                should_unbold = True
            elif "baholash ob" in p_text_lower and "balansidagi" in p_text_lower and ("bozorda yuqori" in p_text_lower or "ikkilamchi bozorda" in p_text_lower):
                should_unbold = True
                
            if should_unbold:
                enforce_paragraph_unbold(p)
                for r in p.xpath('.//w:r'):
                    enforce_arial_normal(r)
                    
            # Force Arial 8, Navy color and right-alignment for company address/bank details block
            is_address_paragraph = False
            if "yuridik manzil" in p_text_lower and ("o'lkansoy" in p_text_lower or "o‘lkansoy" in p_text_lower or "farg" in p_text_lower):
                is_address_paragraph = True
            elif "ofis:" in p_text_lower and ("furqat" in p_text_lower or "farg" in p_text_lower):
                is_address_paragraph = True
            elif "bank:" in p_text_lower and ("universal" in p_text_lower or "mfo" in p_text_lower):
                is_address_paragraph = True
            elif "h/r" in p_text_lower and ("20208" in p_text_lower or "inn" in p_text_lower):
                is_address_paragraph = True
                
            if is_address_paragraph:
                try:
                    p_obj = Paragraph(p, part)
                    p_obj.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    p_obj.paragraph_format.space_before = Pt(0)
                    p_obj.paragraph_format.space_after = Pt(2)
                    p_obj.paragraph_format.line_spacing = 1.0
                    
                    for r in p.xpath('.//w:r'):
                        rPrs = r.xpath('./w:rPr')
                        if not rPrs:
                            rPr = OxmlElement('w:rPr')
                            r.insert(0, rPr)
                        else:
                            rPr = rPrs[0]
                            
                        # Force Arial
                        for f in rPr.xpath('./w:rFonts'):
                            rPr.remove(f)
                        rFonts = OxmlElement('w:rFonts')
                        rFonts.set(qn('w:ascii'), 'Arial')
                        rFonts.set(qn('w:hAnsi'), 'Arial')
                        rFonts.set(qn('w:cs'), 'Arial')
                        rPr.append(rFonts)
                        
                        # Force Size 8 (16 half-points)
                        for sz in rPr.xpath('./w:sz | ./w:szCs'):
                            rPr.remove(sz)
                        sz = OxmlElement('w:sz')
                        sz.set(qn('w:val'), '16')
                        rPr.append(sz)
                        szCs = OxmlElement('w:szCs')
                        szCs.set(qn('w:val'), '16')
                        rPr.append(szCs)
                        
                        # Force Navy Color (002060)
                        for color in rPr.xpath('./w:color'):
                            rPr.remove(color)
                        color = OxmlElement('w:color')
                        color.set(qn('w:val'), '002060')
                        rPr.append(color)
                        
                        # Remove bold and italic
                        for b in rPr.xpath('./w:b | ./w:bCs'):
                            rPr.remove(b)
                        for i in rPr.xpath('./w:i | ./w:iCs'):
                            rPr.remove(i)
                except Exception as addr_err:
                    print(f"Error styling address paragraph: {addr_err}")
                    
        # NOTE: a:p (DrawingML) paragraphs are intentionally NOT processed here.
        # Modifying a:p elements (shapes, charts, images) with process_xml_paragraph
        # corrupts the DrawingML XML relationships and causes Word to report the file as damaged.
        # Text replacements in shapes/charts are handled by docxtpl's Jinja2 rendering above.
    
    # Electronic Signature / Stamp placement
    report_status = context.get('report_status', 'draft')
    is_approved = (report_status == 'approved')
    stamp_date = context.get('{hisobot_sanasi}', datetime.now().strftime('%d.%m.%Y'))
    
    # 1. Place the verification QR code in the right cell of the first page header (top-right corner)
    header_tables = []
    # Add tables from headers
    for section in doc_docx.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header:
                for t in header.tables:
                    if t not in header_tables:
                        header_tables.append(t)
    # Also add tables from document body to handle templates where cover page header is in the body
    for t in doc_docx.tables:
        if t not in header_tables:
            header_tables.append(t)
            
    for table in header_tables:
        for row in table.rows:
            if len(row.cells) >= 2:
                cell_0 = row.cells[0]
                cell_1 = row.cells[1]
                
                text_0 = "".join(p.text for p in cell_0.paragraphs)
                text_1 = "".join(p.text for p in cell_1.paragraphs)
                
                text_0_lower = text_0.lower()
                text_1_lower = text_1.lower()
                
                # Strict matching to ensure we don't accidentally match normal body tables
                keywords_sertifikat = ['сертификат', 'sertifikat', 'hisobot', 'hisobotni', 'hujjat', 'hujjatni']
                keywords_verify = ['амалдалигини', 'amaldaligini', 'haqiqiyligini', 'tekshirish', 'aniqlash', 'амалда', 'amalda', 'tasdiqlash']
                
                has_sertifikat_0 = any(kw in text_0_lower for kw in keywords_sertifikat)
                has_sertifikat_1 = any(kw in text_1_lower for kw in keywords_sertifikat)
                has_verify_0 = any(kw in text_0_lower for kw in keywords_verify)
                has_verify_1 = any(kw in text_1_lower for kw in keywords_verify)
                
                is_verification_table = (has_sertifikat_0 or has_sertifikat_1) and (has_verify_0 or has_verify_1)
                
                if is_verification_table:
                    has_kw_0 = has_sertifikat_0 or has_verify_0
                    # Target Cell 1 (right cell) for the verification text and QR code
                    target_cell = cell_1
                    
                    # Prevent Word from dynamically collapsing Cell 1 by disabling autofit and setting widths
                    table.autofit = False
                    table.allow_autofit = False
                    
                    try:
                        cell_0.width = Inches(4.2)
                        cell_1.width = Inches(2.2)
                        table.columns[0].width = Inches(4.2)
                        table.columns[1].width = Inches(2.2)
                    except Exception as width_err:
                        print(f"Error setting header column/cell widths: {width_err}")
                    
                    verify_text = ""
                    if has_kw_0:
                        # Extract text paragraphs from Cell 0 while preserving paragraphs containing drawings
                        verify_paragraphs = []
                        for p in list(cell_0.paragraphs):
                            has_drawing = False
                            for run in p.runs:
                                if 'w:drawing' in run._r.xml or 'w:pict' in run._r.xml:
                                    has_drawing = True
                            if not has_drawing:
                                if p.text.strip():
                                    verify_paragraphs.append(p.text.strip())
                                p.text = ""
                        verify_text = "\n".join(verify_paragraphs)
                    else:
                        verify_text = text_1.strip()
                        
                    # Clear Cell 1 completely of placeholders or older signature blocks
                    tc_1 = cell_1._tc
                    tcPr_1 = tc_1.get_or_add_tcPr()
                    tc_1.clear()
                    if tcPr_1 is not None:
                        tc_1.append(tcPr_1)
                        
                    # Add verification text to Cell 1 styled beautifully
                    if verify_text:
                        for line in verify_text.split('\n'):
                            if line.strip():
                                p_text = cell_1.add_paragraph()
                                p_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                p_text.paragraph_format.space_before = Pt(0)
                                p_text.paragraph_format.space_after = Pt(2)
                                run_text = p_text.add_run(line.strip())
                                run_text.font.name = 'Arial'
                                run_text.font.size = Pt(8.5)
                                run_text.font.color.rgb = RGBColor(128, 128, 128) # Muted Gray
                                run_text.bold = False
                                
                    # Add QR Code paragraph to Cell 1
                    p_qr = cell_1.add_paragraph()
                    p_qr.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_qr.paragraph_format.space_before = Pt(4)
                    p_qr.paragraph_format.space_after = Pt(2)
                    
                    if qr_code_path and os.path.exists(qr_code_path):
                        verify_qr_file = qr_code_path
                    else:
                        import qrcode
                        verify_url = context.get('verify_absolute_url', 'http://smartbaholash.uz')
                        temp_dir = os.path.dirname(output_path)
                        verify_qr_file = os.path.join(temp_dir, "verify_qr.png")
                        qr = qrcode.QRCode(version=1, box_size=3, border=1)
                        qr.add_data(verify_url)
                        qr.make(fit=True)
                        img = qr.make_image(fill_color="black", back_color="white")
                        img.save(verify_qr_file)
                        
                    try:
                        run_qr = p_qr.add_run()
                        run_qr.add_picture(verify_qr_file, width=Inches(0.65))
                    except Exception as qre:
                        print(f"Error adding verification QR to header: {qre}")
                        
                    # Add unique verification code under QR code
                    try:
                        p_code = cell_1.add_paragraph()
                        p_code.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p_code.paragraph_format.space_before = Pt(2)
                        p_code.paragraph_format.space_after = Pt(2)
                        
                        report_uuid = context.get('report_id_str', '')
                        short_code = report_uuid.split('-')[0].upper() if '-' in report_uuid else report_uuid[:8].upper()
                        if not short_code:
                            short_code = "VERIFIED"
                        
                        run_code = p_code.add_run(f"smartbaholash.uz\nKOD: SV-{short_code}")
                        run_code.font.name = 'Arial'
                        run_code.font.size = Pt(6.5)
                        run_code.font.color.rgb = RGBColor(128, 128, 128)
                        run_code.bold = False
                    except Exception as ce:
                        print(f"Error adding verification code text: {ce}")

    for table in doc_docx.tables:
        # Check if this table is the V.4 table.
        # The table of interest has header 'V.4. Baholovchi tashkilot va baholovchi.' (sometimes just 'V.4.' or 'Baholovchi tashkilot')
        # Let's inspect the cells of the table to see if it's the V.4 table.
        is_v4_table = False
        for r_idx, r in enumerate(table.rows):
            for c_idx, c in enumerate(r.cells):
                c_text = c.text.lower()
                if "baholovchi tashkilot va baholovchi" in c_text or "v.4." in c_text or "baholovchi mutaxassis" in c_text:
                    is_v4_table = True
                    break
            if is_v4_table:
                break
        for row in table.rows:
            if len(row.cells) >= 2:
                combined_text = " ".join(c.text.lower() for c in row.cells)
                is_sig_row = any(kw in combined_text for kw in ['baholovchi', 'rahbari', 'raxbari', 'рахбари', 'раҳбари', 'direktor', 'bajaruvchi', 'tasdiq', 'тасдиq', 'тасдиқ', 'tasdiqlayman', 'тасдиqлайman', 'тасдиqlayman', 'тасдиқлайman', 'тасдиқлайман'])
                
                # Exclude metadata tables containing identity documents/personal details
                is_metadata_row = any(kw in combined_text for kw in ['hujjat', 'pasport', 'jshshir', 'guvohnoma', 'dog\'', 'tel', 'telefon', 'tuman', 'viloyat'])
                if is_sig_row and not is_v4_table and not is_metadata_row:
                    target_cell = row.cells[1]
                    is_appraiser = any(kw in combined_text for kw in ['baholovchi', 'bajaruvchi'])
                    is_tasdiq_block = any(kw in combined_text for kw in ['tasdiq', 'тасдиq', 'тасдиқ', 'tasdiqlayman', 'тасdiqлайman', 'тасдиqlayman', 'тасdiqlayman', 'тасдиқлайман', 'rahbari', 'raxbari', 'рахбари', 'раҳбари', 'direktor'])
                    
                    # If this cell is the verification QR cell itself, skip it here
                    cell_text_lower = "".join(p.text.lower() for p in target_cell.paragraphs)
                    if any(kw in cell_text_lower for kw in ['сертификат', 'амалдалигини', 'aniqlash']):
                        continue
                        
                    cell_text_clean = target_cell.text.strip().replace('_', '').replace(' ', '').replace('\n', '')
                    has_placeholder = any(kw in target_cell.text.lower() for kw in ['tasdiq', 'тасдиқ', 'tasdiqlayman', 'тасдиqлайman', 'тасдиqlayman', 'тасдиқлайман'])
                    
                    if not cell_text_clean or len(cell_text_clean) < 3 or has_placeholder:
                        # Clear target cell completely of all text/images/runs
                        tc = target_cell._tc
                        tcPr = tc.get_or_add_tcPr()
                        tc.clear()
                        if tcPr is not None:
                            tc.append(tcPr)
                              
                        # Draw premium E-signature border
                        from docx.oxml import parse_xml
                        from docx.oxml.ns import nsdecls
                        
                        if is_approved:
                            border_color = "22C55E" # Green-500
                            bg_color = "F0FDF4"     # Light Green (Green-50)
                        else:
                            border_color = "EF4444" # Red-500
                            bg_color = "FEF2F2"     # Light Red (Red-50)
                            
                        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color}"/>')
                        tcPr.append(shd)
                        borders = parse_xml(f'''
                            <w:tcBorders {nsdecls("w")}>
                                <w:top w:val="single" w:sz="6" w:space="0" w:color="{border_color}"/>
                                <w:left w:val="single" w:sz="6" w:space="0" w:color="{border_color}"/>
                                <w:bottom w:val="single" w:sz="6" w:space="0" w:color="{border_color}"/>
                                <w:right w:val="single" w:sz="6" w:space="0" w:color="{border_color}"/>
                            </w:tcBorders>
                        ''')
                        tcPr.append(borders)
                        
                        # Add paragraph for E-signature details
                        p_details = target_cell.add_paragraph()
                        p_details.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p_details.paragraph_format.space_before = Pt(2)
                        p_details.paragraph_format.space_after = Pt(2)
                        p_details.paragraph_format.line_spacing = 1.0
                        
                        if is_approved:
                            run = p_details.add_run("● ELEKTRON IMZO BILAN TASDIQLANGAN\n")
                            run.font.name = 'Arial'
                            run.font.size = Pt(6.5)
                            run.font.color.rgb = RGBColor(22, 163, 74) # Green-600
                            run.bold = True
                            
                            appraiser_name = context.get('appraiser_name', 'B.M.Mirzabekov')
                            appraiser_license = context.get('appraiser_license', '№0244')
                            company_name = context.get('company_name', 'PERCEPTION VALUE')
                            company_stir = context.get('company_stir', '301234567')
 
                            cert_text = f"Sertifikat: {appraiser_license}" if is_appraiser else f"STIR: {company_stir}"
                            run2 = p_details.add_run(f"Tizim: smartbaholash.uz\n{cert_text}\nSana: {stamp_date}")
                            run2.font.name = 'Arial'
                            run2.font.size = Pt(6)
                            run2.font.color.rgb = RGBColor(22, 163, 74)
                        else:
                            run = p_details.add_run("▲ TO'LOV AMALGA OSHIRILMAGAN\n")
                            run.font.name = 'Arial'
                            run.font.size = Pt(6.5)
                            run.font.color.rgb = RGBColor(220, 38, 38) # Red-600
                            run.bold = True
                            
                            run2 = p_details.add_run("Hisobot tasdiqlanmagan!\nTizim: smartbaholash.uz")
                            run2.font.name = 'Arial'
                            run2.font.size = Pt(6)
                            run2.font.color.rgb = RGBColor(220, 38, 38)
                            
                        # Insert signature or seal image inside the cell
                        if is_approved:
                            p_img = target_cell.add_paragraph()
                            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            p_img.paragraph_format.space_before = Pt(2)
                            p_img.paragraph_format.space_after = Pt(2)
                            run_img = p_img.add_run()
                            
                            if is_appraiser:
                                img_file = context.get('signature_path')
                                if img_file and os.path.exists(img_file):
                                    try:
                                        run_img.add_picture(img_file, width=Inches(0.85))
                                    except Exception as e:
                                        print(f"Error adding appraiser signature image: {e}")
                            else:
                                img_file = context.get('seal_path')
                                if img_file and os.path.exists(img_file):
                                    try:
                                        run_img.add_picture(img_file, width=Inches(0.95))
                                    except Exception as e:
                                        print(f"Error adding company seal image: {e}")
                                
                        # If row has 2 cells, populate the name (only if not tasdiq block)
                        if not is_tasdiq_block:
                            dynamic_name = context.get('appraiser_name', 'B.M.Mirzabekov') if is_appraiser else "B.M.Mirzabekov"
                            if len(row.cells) == 2:
                                row.cells[0].width = Inches(3.5)
                                row.cells[1].width = Inches(3.0)
                                p_name = target_cell.add_paragraph()
                                p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                p_name.paragraph_format.space_before = Pt(4)
                                p_name.paragraph_format.space_after = Pt(4)
                                
                                run_name = p_name.add_run(dynamic_name)
                                run_name.font.name = 'Arial'
                                run_name.font.size = Pt(10)
                                run_name.font.color.rgb = RGBColor(0, 0, 128) # Navy Blue
                                run_name.bold = True
                            elif len(row.cells) >= 3:
                                row.cells[0].width = Inches(3.0)
                                row.cells[1].width = Inches(2.0)
                                row.cells[2].width = Inches(2.0)
                                name_cell = row.cells[2]
                                if not name_cell.text.strip():
                                    for np in name_cell.paragraphs:
                                        np.text = ""
                                    np = name_cell.paragraphs[0] if name_cell.paragraphs else name_cell.add_paragraph()
                                    np.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                                    
                                    run_name = np.add_run(dynamic_name)
                                    run_name.font.name = 'Arial'
                                    run_name.font.size = Pt(11)
                                    run_name.font.color.rgb = RGBColor(0, 0, 128) # Navy Blue
                                    run_name.bold = True
                                
    # 3. Replace legacy single-brace QR code tags like {qr_code}, {verify_qr}, {qr} inside the document parts
    if qr_code_path and os.path.exists(qr_code_path):
        qr_tags = ['{qr_code}', '{qr}', '{verify_qr}', '{verification_qr}']
        
        # Helper to search and replace tag with image in a paragraph's runs
        def replace_tags_in_p(p):
            for tag in qr_tags:
                if tag in p.text:
                    # Find the run containing the tag and replace it with image
                    for run in p.runs:
                        if tag in run.text:
                            run.text = run.text.replace(tag, '')
                            run.add_picture(qr_code_path, width=Inches(0.95))
                            
        # Scan in body paragraphs
        for p in doc_docx.paragraphs:
            replace_tags_in_p(p)
            
        # Scan in tables
        for table in doc_docx.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        replace_tags_in_p(p)
                        
        # Scan in headers/footers
        for section in doc_docx.sections:
            for part_name in ['header', 'first_page_header', 'even_page_header', 'footer', 'first_page_footer', 'even_page_footer']:
                part = getattr(section, part_name, None)
                if part:
                    for p in part.paragraphs:
                        replace_tags_in_p(p)
                    for table in part.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for p in cell.paragraphs:
                                    replace_tags_in_p(p)

    # 4. Replace legacy single-brace Signature/Seal tags like {signature}, {seal}, {imzo}, {pechat} inside the document parts
    sig_tags = ['{signature}', '{appraiser_signature}', '{sig}', '{imzo}']
    seal_tags = ['{seal}', '{company_seal}', '{stamp}', '{pechat}']
    sig_file = context.get('signature_path') if is_approved else None
    seal_file = context.get('seal_path') if is_approved else None
    
    def replace_sig_seal_tags_in_p(p):
        for tag in sig_tags:
            if tag in p.text:
                for run in p.runs:
                    if tag in run.text:
                        run.text = run.text.replace(tag, '')
                        if sig_file and os.path.exists(sig_file):
                            run.add_picture(sig_file, width=Inches(0.85))
        for tag in seal_tags:
            if tag in p.text:
                for run in p.runs:
                    if tag in run.text:
                        run.text = run.text.replace(tag, '')
                        if seal_file and os.path.exists(seal_file):
                            run.add_picture(seal_file, width=Inches(0.95))
                            
        # Scan in body paragraphs
        for p in doc_docx.paragraphs:
            replace_sig_seal_tags_in_p(p)
            
        # Scan in tables
        for table in doc_docx.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        replace_sig_seal_tags_in_p(p)
                        
        # Scan in headers/footers
        for section in doc_docx.sections:
            for part_name in ['header', 'first_page_header', 'even_page_header', 'footer', 'first_page_footer', 'even_page_footer']:
                part = getattr(section, part_name, None)
                if part:
                    for p in part.paragraphs:
                        replace_sig_seal_tags_in_p(p)
                    for table in part.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for p in cell.paragraphs:
                                    replace_sig_seal_tags_in_p(p)

    doc_docx.save(output_path)
    return output_path
    
    # 5. Replace legacy single-brace combined E-Signature QR block tags inside the document parts
    combined_tags = ['{imzo_qr}', '{signature_qr}', '{approve_qr}', '{verification_block}', '{verify_block}']
    
    # Search inside tables to style cells with a green/red-bordered box
    for table in doc_docx.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in list(cell.paragraphs):
                    for tag in combined_tags:
                        if tag in p.text:
                            tc = cell._tc
                            tcPr = tc.get_or_add_tcPr()
                            tc.clear()
                            if tcPr is not None:
                                tc.append(tcPr)
                                
                            from docx.oxml import parse_xml
                            from docx.oxml.ns import nsdecls
                            
                            if is_approved:
                                border_color = "22C55E" # Green-500
                                bg_color = "F0FDF4"     # Light Green (Green-50)
                            else:
                                border_color = "EF4444" # Red-500
                                bg_color = "FEF2F2"     # Light Red (Red-50)
                            
                            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color}"/>')
                            tcPr.append(shd)
                            borders = parse_xml(f'''
                                <w:tcBorders {nsdecls("w")}>
                                    <w:top w:val="single" w:sz="6" w:space="0" w:color="{border_color}"/>
                                    <w:left w:val="single" w:sz="6" w:space="0" w:color="{border_color}"/>
                                    <w:bottom w:val="single" w:sz="6" w:space="0" w:color="{border_color}"/>
                                    <w:right w:val="single" w:sz="6" w:space="0" w:color="{border_color}"/>
                                </w:tcBorders>
                            ''')
                            tcPr.append(borders)
                            
                            if is_approved:
                                # Insert QR Code inside the styled cell
                                p_qr = cell.add_paragraph()
                                p_qr.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                p_qr.paragraph_format.space_before = Pt(4)
                                p_qr.paragraph_format.space_after = Pt(2)
                                run_qr = p_qr.add_run()
                                run_qr.add_picture(verify_qr_file, width=Inches(0.95))
                            
                            # Insert description text
                            p_text = cell.add_paragraph()
                            p_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            p_text.paragraph_format.space_before = Pt(2)
                            p_text.paragraph_format.space_after = Pt(4)
                            p_text.paragraph_format.line_spacing = 1.0
                            
                            if is_approved:
                                run_head = p_text.add_run("● ELEKTRON IMZO BILAN TASDIQLANGAN\n")
                                run_head.font.name = 'Arial'
                                run_head.font.size = Pt(6.5)
                                run_head.font.color.rgb = RGBColor(22, 163, 74)
                                run_head.bold = True
                                
                                appraiser_license = context.get('appraiser_license', '№0244')
                                run_body = p_text.add_run(f"Tizim: smartbaholash.uz\nSertifikat: {appraiser_license}\nSana: {stamp_date}")
                                run_body.font.name = 'Arial'
                                run_body.font.size = Pt(6)
                                run_body.font.color.rgb = RGBColor(22, 163, 74)
                            else:
                                run_head = p_text.add_run("▲ TO'LOV AMALGA OSHIRILMAGAN\n")
                                run_head.font.name = 'Arial'
                                run_head.font.size = Pt(6.5)
                                run_head.font.color.rgb = RGBColor(220, 38, 38)
                                run_head.bold = True
                                
                                run_body = p_text.add_run("Hisobot tasdiqlanmagan!\nTizim: smartbaholash.uz")
                                run_body.font.name = 'Arial'
                                run_body.font.size = Pt(6)
                                run_body.font.color.rgb = RGBColor(220, 38, 38)
                            break
                            
    # Search inside normal paragraphs outside tables
    for p in doc_docx.paragraphs:
        for tag in combined_tags:
            if tag in p.text:
                p.text = p.text.replace(tag, '')
                if is_approved:
                    run_qr = p.add_run()
                    run_qr.add_picture(verify_qr_file, width=Inches(0.95))

                p_next = doc_docx.add_paragraph()
                p_next.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if is_approved:
                    run_head = p_next.add_run("● ELEKTRON IMZO BILAN TASDIQLANGAN\n")
                    run_head.font.name = 'Arial'
                    run_head.font.size = Pt(6.5)
                    run_head.font.color.rgb = RGBColor(22, 163, 74)
                    run_head.bold = True
                    
                    appraiser_license = context.get('appraiser_license', '№0244')
                    run_body = p_next.add_run(f"Tizim: smartbaholash.uz\nSertifikat: {appraiser_license}\nSana: {stamp_date}")
                    run_body.font.name = 'Arial'
                    run_body.font.size = Pt(6)
                    run_body.font.color.rgb = RGBColor(22, 163, 74)
                else:
                    run_head = p_next.add_run("▲ TO'LOV AMALGA OSHIRILMAGAN\n")
                    run_head.font.name = 'Arial'
                    run_head.font.size = Pt(6.5)
                    run_head.font.color.rgb = RGBColor(220, 38, 38)
                    run_head.bold = True
                    
                    run_body = p_next.add_run("Hisobot tasdiqlanmagan!\nTizim: smartbaholash.uz")
                    run_body.font.name = 'Arial'
                    run_body.font.size = Pt(6)
                    run_body.font.color.rgb = RGBColor(220, 38, 38)

    doc_docx.save(output_path)
    return output_path

def generate_professional_table(doc, data, insert_after=None):
    table = doc.add_table(rows=1, cols=5)
    if insert_after:
        insert_after._p.addnext(table._element)
    table.style = 'Table Grid'
    table.autofit = True
    
    # Headers
    hdr_cells = table.rows[0].cells
    headers = ['Asosiy ko‘rsatkichlar', 'Baholash ob’yekti', 'Analog I', 'Analog II', 'Analog III']
    for i, h_text in enumerate(headers):
        hdr_cells[i].text = h_text
        set_cell_arial_12(hdr_cells[i], bold=True)
    
    # Professional rows mapping from context data
    for row_key, row_vals in data.get('rows', []):
        row = table.add_row().cells
        row[0].text = str(row_key)
        set_cell_arial_12(row[0], bold=True)
        
        for i, val in enumerate(row_vals):
            if i + 1 < 5:
                # If it's the last row (Manba) and value is a URL, use a hyperlink
                if row_key == "manba" and str(val).startswith('http'):
                    row[i+1].text = ""  # Clear default text
                    par = row[i+1].paragraphs[0]
                    add_hyperlink(par, str(val), "Havola")
                    set_cell_arial_12(row[i+1])
                else:
                    row[i+1].text = str(val if val is not None else "-")
                    set_cell_arial_12(row[i+1])

def generate_criteria_table(doc, data, insert_after=None):
    table = doc.add_table(rows=1, cols=4)
    if insert_after:
        insert_after._p.addnext(table._element)
    table.style = 'Table Grid'
    table.autofit = True
    
    hdr = table.rows[0].cells
    headers = ['№', 'Mezonlar', 'Qiyosiy', 'Xarajat']
    for i, h_text in enumerate(headers):
        hdr[i].text = h_text
        set_cell_arial_12(hdr[i], bold=True)
    
    for i, (name, comp_val, cost_val) in enumerate(data.get('rows', []), 1):
        cells = table.add_row().cells
        cells[0].text = str(i)
        cells[1].text = str(name)
        cells[2].text = str(comp_val)
        cells[3].text = str(cost_val)
        for c in cells:
            set_cell_arial_12(c)

def generate_reconciliation_table(doc, data, insert_after=None):
    table = doc.add_table(rows=1, cols=4)
    if insert_after:
        insert_after._p.addnext(table._element)
    table.style = 'Table Grid'
    table.autofit = True
    
    hdr = table.rows[0].cells
    headers = ['№', 'Omillar', 'Xarajat', 'Qiyosiy']
    for i, h_text in enumerate(headers):
        hdr[i].text = h_text
        set_cell_arial_12(hdr[i], bold=True)
    
    for i, (name, cost_val, comp_val) in enumerate(data.get('rows', []), 1):
        cells = table.add_row().cells
        cells[0].text = str(i)
        cells[1].text = str(name)
        cells[2].text = str(cost_val)
        cells[3].text = str(comp_val)
        for c in cells:
            set_cell_arial_12(c)

def generate_macro_table(doc, data, insert_after=None):
    """
    Generates Section VI.1 table for macroeconomic indicators.
    """
    table = doc.add_table(rows=1, cols=4)
    if insert_after:
        insert_after._p.addnext(table._element)
    table.style = 'Table Grid'
    table.autofit = True
    
    hdr = table.rows[0].cells
    headers = [
        'Makroiqtisodiy Ko‘rsatkich', 
        data.get('period_label', '2025 yil yanvar-iyun'), 
        'O‘sish Sur’ati', 
        'Avtomobil Bozoriga Ta’siri'
    ]
    
    # Header styling
    for i, h_text in enumerate(headers):
        hdr[i].text = h_text
        set_cell_arial_12(hdr[i], bold=True)

    for row_data in data.get('rows', []):
        cells = table.add_row().cells
        for i, val in enumerate(row_data):
            cells[i].text = str(val)
            set_cell_arial_12(cells[i])
