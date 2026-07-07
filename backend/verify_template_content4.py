import os
import django
import docx

os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'valuation_platform.settings')
django.setup()

from reports.models import ReportTemplate

def analyze_template_content():
    templates = ReportTemplate.objects.all().order_by('-created_at')[:1]
    with open('template_dump4.txt', 'w', encoding='utf-8') as f:
        f.write("Searching for Dvigatel, Kuzov, Shassi in the latest template...\n\n")
        
        for template in templates:
            try:
                doc = docx.Document(template.file.path)
                
                for p in doc.paragraphs:
                    text_lower = p.text.lower()
                    if 'dvigatel' in text_lower or 'kuzov' in text_lower or 'shassi' in text_lower or 'ishlab chiqarilgan' in text_lower:
                        f.write(f"FOUND: {p.text}\n")
                        
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text_lower = cell.text.lower()
                            if 'dvigatel' in text_lower or 'kuzov' in text_lower or 'shassi' in text_lower or 'ishlab chiqarilgan' in text_lower:
                                f.write(f"TABLE FOUND: {cell.text.replace(chr(10), ' ')}\n")
                                
            except Exception as e:
                f.write(f"Error reading docx: {e}\n")

if __name__ == '__main__':
    analyze_template_content()
