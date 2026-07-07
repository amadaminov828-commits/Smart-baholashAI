from docx import Document
import os

def create_sample_smart_template():
    target_path = 'c:/Users/Asus/Desktop/antigravity/backend/media/templates/smart_re_template.docx'
    os.makedirs(os.path.dirname(target_path), exist_ok=True)
    
    doc = Document()
    doc.add_heading('Smart Baholash: Ko`chmas Mulk (Aqlli Shabloni)', 0)
    
    doc.add_paragraph('Ushbu shablon avtomatik ravishda tanlangan yo`nalishga moslashadi.')
    
    doc.add_heading('Umumiy ma`lumotlar', level=1)
    doc.add_paragraph('Mijoz: {{ FISH }}')
    doc.add_paragraph('Kadastr raqami: {{ KADASTR }}')
    doc.add_paragraph('Yo`nalish: {{ PURPOSE }}')
    doc.add_paragraph('Sana: {{ bugun }}')
    
    # Conditional logic examples
    doc.add_heading('Maxsus bo`limlar', level=1)
    
    p1 = doc.add_paragraph()
    p1.add_run('{% if is_garov %}').bold = True
    doc.add_paragraph('🛑 BU BO`LIM FAQAT GAROV UCHUN: Bank uchun maxsus shartlar...')
    p1_end = doc.add_paragraph()
    p1_end.add_run('{% endif %}').bold = True
    
    p2 = doc.add_paragraph()
    p2.add_run('{% if is_soliq %}').bold = True
    doc.add_paragraph('💰 BU BO`LIM FAQAT SOLIQ UCHUN: Soliq bazasini hisoblash usuli...')
    p2_end = doc.add_paragraph()
    p2_end.add_run('{% endif %}').bold = True
    
    p3 = doc.add_paragraph()
    p3.add_run('{% if is_snos %}').bold = True
    doc.add_paragraph('🏗️ BU BO`LIM FAQAT SNOS UCHUN: Kompensatsiya miqdori va zarar hisobi...')
    p3_end = doc.add_paragraph()
    p3_end.add_run('{% endif %}').bold = True
    
    doc.add_heading('QR Kod', level=1)
    doc.add_paragraph('{{ qr_code }}')
    
    doc.save(target_path)
    print(f"Sample template created at: {target_path}")

if __name__ == "__main__":
    create_sample_smart_template()
