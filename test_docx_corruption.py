"""
Diagnostic script to find which step in docx_generator.py corrupts the DOCX.
Saves intermediate files at each stage and tests each one with Word COM.
Runs directly on the existing corrupted report_539.docx's template.
"""
import os, sys, comtypes.client

# Work from backend directory
BASE_DIR = os.path.dirname(__file__)
BACKEND_DIR = os.path.join(BASE_DIR, 'backend')
sys.path.insert(0, BACKEND_DIR)
os.chdir(BACKEND_DIR)

# Find the template file and existing docx
import glob
templates = glob.glob('media/templates/*.docx') + glob.glob('media/**/*.docx', recursive=True)
templates = [t for t in templates if 'report_' not in t]
print(f"Templates found: {templates}")

# Find latest report docx
docx_files = glob.glob('media/**/*.docx', recursive=True)
docx_files.sort(key=os.path.getmtime, reverse=True)
print(f"Docx files found: {docx_files[:3]}")

# Use known existing report docx for stage tests
EXISTING_DOCX = os.path.abspath('media/tmp/88e60f7bb7f4/report_539.docx')

SCRATCH_DIR = os.path.join(BASE_DIR, 'backend', 'debug_stages')
os.makedirs(SCRATCH_DIR, exist_ok=True)

def test_docx_with_word(path, label):
    """Try to open and convert the DOCX file with Word. Returns True if OK."""
    word = comtypes.client.CreateObject('Word.Application')
    word.Visible = False
    pdf_out = path.replace('.docx', '_test.pdf')
    try:
        doc = word.Documents.Open(os.path.abspath(path))
        doc.SaveAs(os.path.abspath(pdf_out), FileFormat=17)
        doc.Close()
        word.Quit()
        print(f"  ✓ {label}: VALID (PDF OK)")
        return True
    except Exception as e:
        try:
            word.Quit()
        except:
            pass
        print(f"  ✗ {label}: CORRUPTED - {e}")
        return False

from docxtpl import DocxTemplate
from docx.shared import Inches, Pt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
import datetime

# ---- STAGE 0: Test the existing corrupted file ----
print("\n=== STAGE 0: Existing DOCX file ===")
if os.path.exists(EXISTING_DOCX):
    print(f"  File: {EXISTING_DOCX}")
    test_docx_with_word(EXISTING_DOCX, "Stage 0 (existing docx)")
else:
    print(f"  File not found: {EXISTING_DOCX}")

# Find template
TEMPLATE_PATH = None
for t in templates:
    if os.path.exists(t):
        TEMPLATE_PATH = os.path.abspath(t)
        break
if TEMPLATE_PATH is None:
    # Try to find by name hint
    all_tpl = glob.glob('**/*.docx', recursive=True)
    for f in all_tpl:
        if 'template' in f.lower() or 'shablon' in f.lower():
            TEMPLATE_PATH = os.path.abspath(f)
            break

print(f"\nTemplate: {TEMPLATE_PATH}")
if not TEMPLATE_PATH:
    print("ERROR: No template found - need to use template file")
    sys.exit(1)

# Build minimal context
context = {'report_id_str': 'TEST-001', 'hisobot_sanasi': datetime.date.today().strftime('%d.%m.%Y'), 'report_status': 'draft'}

# ---- STAGE 1: Just render with docxtpl ----
print("\n=== STAGE 1: docxtpl render only ===")
stage1_path = os.path.join(SCRATCH_DIR, 'stage1_jinja_only.docx')
try:
    tpl = DocxTemplate(TEMPLATE_PATH)
    tpl.render(context, autoescape=True)
    tpl.save(stage1_path)
    print(f"  Saved: {stage1_path}")
    test_docx_with_word(stage1_path, "Stage 1 (jinja only)")
except Exception as e:
    import traceback
    traceback.print_exc()
    print(f"  FAILED: {e}")
    stage1_path = None

# ---- STAGE 2: After python-docx reload ----
print("\n=== STAGE 2: After python-docx reload only ===")
stage2_path = os.path.join(SCRATCH_DIR, 'stage2_reload.docx')
try:
    from docx import Document
    doc = Document(stage1_path)
    doc.save(stage2_path)
    test_docx_with_word(stage2_path, "Stage 2 (reload + save)")
except Exception as e:
    print(f"  FAILED: {e}")

# ---- STAGE 3: After process_xml_paragraph loop ----
print("\n=== STAGE 3: After process_xml_paragraph loop ===")
stage3_path = os.path.join(SCRATCH_DIR, 'stage3_after_paragraph_loop.docx')
try:
    from vehicles.docx_generator import process_xml_paragraph, enforce_arial_normal, enforce_paragraph_unbold, populate_formatted_text
    from docx.oxml.ns import qn
    from docx.shared import RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import lxml.etree as etree
    
    doc = Document(stage1_path)
    
    # Table autofit
    for table in doc.tables:
        table.autofit = False
    
    # Run process_xml_paragraph on all body paragraphs
    for p in doc.paragraphs:
        try:
            process_xml_paragraph(p, context, doc)
        except:
            pass
    
    # Run on table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    try:
                        process_xml_paragraph(p, context, doc)
                    except:
                        pass
    
    doc.save(stage3_path)
    test_docx_with_word(stage3_path, "Stage 3 (after process_xml_paragraph)")
except Exception as e:
    import traceback
    print(f"  FAILED: {e}")
    traceback.print_exc()

# ---- STAGE 4: After header table / QR code placement ----
print("\n=== STAGE 4: After header QR code placement ===")
stage4_path = os.path.join(SCRATCH_DIR, 'stage4_after_header_qr.docx')
try:
    doc = Document(stage1_path)
    
    import qrcode, io
    from PIL import Image
    
    verify_qr_file = os.path.join(SCRATCH_DIR, 'test_qr.png')
    qr = qrcode.QRCode(version=1, box_size=3, border=1)
    qr.add_data('http://smartbaholash.uz/test')
    qr.make(fit=True)
    img = qr.make_image(fill_color='black', back_color='white')
    img.save(verify_qr_file)
    
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    
    # The header table QR code insertion (same as in docx_generator lines 776-911)
    keywords_sertifikat = ['сертификат', 'sertifikat', 'hisobot', 'hisobotni', 'hujjat', 'hujjatni']
    keywords_verify = ['амалдалигини', 'amaldaligini', 'haqiqiyligini', 'tekshirish', 'aniqlash', 'амалда', 'amalda', 'tasdiqlash']
    
    header_tables = []
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header:
                for t in header.tables:
                    if t not in header_tables:
                        header_tables.append(t)
    for t in doc.tables:
        if t not in header_tables:
            header_tables.append(t)
    
    qr_placed = False
    for table in header_tables:
        for row in table.rows:
            if len(row.cells) >= 2:
                cell_0 = row.cells[0]
                cell_1 = row.cells[1]
                text_0 = ''.join(p.text for p in cell_0.paragraphs).lower()
                text_1 = ''.join(p.text for p in cell_1.paragraphs).lower()
                
                has_sertifikat_0 = any(kw in text_0 for kw in keywords_sertifikat)
                has_sertifikat_1 = any(kw in text_1 for kw in keywords_sertifikat)
                has_verify_0 = any(kw in text_0 for kw in keywords_verify)
                has_verify_1 = any(kw in text_1 for kw in keywords_verify)
                
                if (has_sertifikat_0 or has_sertifikat_1) and (has_verify_0 or has_verify_1):
                    print(f"  Found verification table, placing QR code...")
                    qr_placed = True
                    # Place QR in cell_1
                    tc_1 = cell_1._tc
                    tcPr_1 = tc_1.get_or_add_tcPr()
                    tc_1.clear()
                    if tcPr_1 is not None:
                        tc_1.append(tcPr_1)
                    
                    p_qr = cell_1.add_paragraph()
                    p_qr.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run_qr = p_qr.add_run()
                    run_qr.add_picture(verify_qr_file, width=Inches(0.65))
                    break
    
    if not qr_placed:
        print("  No verification table found (skipped header QR)")
    
    doc.save(stage4_path)
    test_docx_with_word(stage4_path, "Stage 4 (after header QR)")
except Exception as e:
    import traceback
    print(f"  FAILED: {e}")
    traceback.print_exc()

# ---- STAGE 5: After V4 table sig/seal placement ----
print("\n=== STAGE 5: After V4 table sig/seal placement ===")
stage5_path = os.path.join(SCRATCH_DIR, 'stage5_after_sig_seal.docx')
try:
    doc = Document(stage1_path)
    
    for table in doc.tables:
        is_v4_table = False
        for r in table.rows:
            for c in r.cells:
                if "baholovchi tashkilot va baholovchi" in c.text.lower() or "v.4." in c.text.lower():
                    is_v4_table = True
                    break
            if is_v4_table:
                break
        
        for row in table.rows:
            if len(row.cells) >= 2:
                combined_text = " ".join(c.text.lower() for c in row.cells)
                is_sig_row = any(kw in combined_text for kw in ['baholovchi', 'rahbari', 'direktor', 'tasdiq'])
                is_metadata_row = any(kw in combined_text for kw in ['hujjat', 'pasport', 'jshshir', 'guvohnoma'])
                
                if is_sig_row and not is_v4_table and not is_metadata_row:
                    target_cell = row.cells[1]
                    cell_text_clean = target_cell.text.strip().replace('_','').replace(' ','').replace('\n','')
                    if not cell_text_clean or len(cell_text_clean) < 3:
                        print(f"  Found sig row, populating cell...")
                        tc = target_cell._tc
                        tcPr = tc.get_or_add_tcPr()
                        tc.clear()
                        if tcPr is not None:
                            tc.append(tcPr)
                        p = target_cell.add_paragraph()
                        p.add_run("TEST SIGNATURE")
    
    doc.save(stage5_path)
    test_docx_with_word(stage5_path, "Stage 5 (after sig/seal)")
except Exception as e:
    import traceback
    print(f"  FAILED: {e}")
    traceback.print_exc()

print("\n=== DONE ===")
print(f"Check files in: {SCRATCH_DIR}")
