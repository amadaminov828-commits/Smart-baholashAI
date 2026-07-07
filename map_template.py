import os
import re
from docx import Document

def normalize_text(t):
    t = t.replace('\xa0', ' ')
    t = t.replace('\r', '\n')
    t = t.replace('“', '"').replace('”', '"')
    t = t.replace('‘', "'").replace('’', "'")
    return t

def replace_in_doc(doc_path, out_path):
    doc = Document(doc_path)
    
    replacements = {
        # Dates & Time
        "2025 yil oktyabr oyining 21chi kuni.": "{baholash_sanasi}.",
        "2025 yil oktyabr oyining 21chi kuni": "{baholash_sanasi}",
        "21/10/2025": "{baholash_sanasi_qisqa}",
        
        "2025 yil oktyabr oyining 23-kuni": "{hisobot_sanasi}",
        "2025 yil oktyabr oyining 23chi kuni.": "{hisobot_sanasi}.",
        "2025 yil oktyabr oyining 23chi kuni": "{hisobot_sanasi}",
        "2025 yil “23” oktyabr": "{hisobot_sanasi_qisqa}",
        "2025 yil \"23\" oktyabr": "{hisobot_sanasi_qisqa}",
        
        "191/01": "{hisobot_raqami}",
        "15.10.2025": "{shartnoma_sanasi}",
        "K1095331": "{shartnoma_raqami}",
        
        "40 855 GBA": "{davlat_raqami}",
        "2018": "{ishlab_chiqarilgan_yili}",
        "«Lacetti»": "«{modeli}»",
        "Lacetti": "{modeli}",
        
        "AAF0458563": "{tex_pasport_seriyasi}",
        "B15D211183566DYUX0327": "{dvigatel_raqami}",
        "XWB5V312DKA528430": "{kuzov_raqami}",
        "398 604": "{yurgan_masofasi}",
        
        "«MAXAM CHIRCHIQ» AJ": "{buyurtmachi}",
        "MAXAM CHIRCHIQ": "{buyurtmachi}",
        
        "76 972 000": "{yakuniy_narx}",
        "Yeltmish olti million to'qqiz yuz yetmish ikki ming so'm": "{yakuniy_narx_suzlarda}",
        "Yeltmish olti million to‘qqiz yuz yetmish ikki ming so‘m": "{yakuniy_narx_suzlarda}",
        
        # Vehicle Specs (from Texpasport or Manual)
        "Yengil avtomobil (Sedan, S-klass)": "{transport_turi}",
        "Yengil avtomobil": "{transport_turi}",
        "Benzin / Metan (Gaz uskunasi o‘rnatilgan)": "{yoqilgi_turi}",
        "Benzin / Metan (Gaz uskunasi o'rnatilgan)": "{yoqilgi_turi}",
        "Oq": "{rangi}",
        "Qoniqarli": "{texnik_holati}",
        "128 302 457,42 so‘m": "{dastlabki_balans_qiymati}",
        "128 302 457.42 so'm": "{dastlabki_balans_qiymati}",
        "0,00 so‘m": "{qoldiq_balans_qiymati}",
        "0.00 so'm": "{qoldiq_balans_qiymati}",
        
        "1498 sm³ (1.5 litr)": "{dvigatel_hajmi}",
        "1498 sm³": "{dvigatel_hajmi}",
        "105 ot kuchi": "{dvigatel_quvvati}",
        "Mexanika (MT-5)": "{uzatma_qutisi}",
        "180 km/soat": "{maksimal_tezlik}",
        "1680 kg": "{tola_vazni}",
        "1210 kg": "{yuksiz_vazni}",
        
        # CBU API Rates (USD, EUR, RUB)
        "12 060,15": "{usd_kursi}",
        "12 060.15": "{usd_kursi}",
        "14 064,55": "{eur_kursi}",
        "14 064.55": "{eur_kursi}",
        "148,87": "{rub_kursi}",
        "148.87": "{rub_kursi}"
    }

    sorted_reps = sorted(replacements.items(), key=lambda x: len(x[0]), reverse=True)

    def process_element(p):
        text = p.text
        if not text.strip(): return
        
        changed = False
        n_text = normalize_text(text)
        
        for old, new in sorted_reps:
            n_old = normalize_text(old)
            if old in text:
                text = text.replace(old, new)
                changed = True
            elif n_old in n_text:
                n_text = n_text.replace(n_old, new)
                text = n_text # Fallback
                changed = True
                
        if changed:
            for r in p.runs: r.text = ""
            if p.runs: p.runs[0].text = text
            else: p.add_run(text)

    for p in doc.paragraphs: process_element(p)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs: process_element(p)
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header:
                for p in header.paragraphs: process_element(p)
                for t in header.tables:
                    for row in t.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs: process_element(p)
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer:
                for p in footer.paragraphs: process_element(p)
                for t in footer.tables:
                    for row in t.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs: process_element(p)

    doc.save(out_path)
    print(f"Saved to {out_path}")

if __name__ == '__main__':
    replace_in_doc('Professional_shablon.docx', 'backend/media/templates/professional_template.docx')
