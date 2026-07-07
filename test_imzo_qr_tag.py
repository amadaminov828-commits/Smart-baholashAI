import os
import sys
import docx

sys.path.append(r'c:\Users\Asus\Desktop\antigravity\backend')
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'valuation_platform.settings')
import django
django.setup()

from vehicles.docx_generator import fill_docx_template

# Create a dummy template docx with combined tags
doc = docx.Document()
doc.add_paragraph("Salom, bu test hisobot.")
doc.add_paragraph("Ulanma tagi: {imzo_qr}")

# Add a table with placeholder
table = doc.add_table(rows=1, cols=2)
table.rows[0].cells[0].text = "Imzolar:"
table.rows[0].cells[1].text = "{signature_qr}"

template_path = "test_imzo_qr_template.docx"
output_path = "test_imzo_qr_output.docx"

doc.save(template_path)

context = {
    "report_id_str": "12345678-abcd-1234-abcd-1234567890ab",
    "report_status": "approved",
    "appraiser_license": "№0999",
    "verify_absolute_url": "http://smartbaholash.uz/verify/123456"
}
qr_image_path = r"c:\Users\Asus\Desktop\antigravity\backend\digital_sig.png"

fill_docx_template(template_path, output_path, context, qr_code_path=qr_image_path)

# Verify placeholders are replaced and images are present
out_doc = docx.Document(output_path)
print("=== Output Document Content ===")
for p in out_doc.paragraphs:
    print(f"P: {p.text}")
    for run in p.runs:
        if 'w:drawing' in run._r.xml or 'w:pict' in run._r.xml:
            print("  [Contains Image]")

for t in out_doc.tables:
    for r in t.rows:
        row_text = [cell.text for cell in r.cells]
        print(f"Table Row: {row_text}")
        for cell in r.cells:
            # Check cell background
            tcPr = cell._tc.get_or_add_tcPr()
            shd = tcPr.xpath('./w:shd')
            if shd:
                print(f"  Cell background: {shd[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')}")
            borders = tcPr.xpath('./w:tcBorders')
            if borders:
                print("  Cell borders present")

# Clean up
if os.path.exists(template_path):
    os.remove(template_path)
if os.path.exists(output_path):
    os.remove(output_path)
